"""Tests for the DOCX parser.

Creates minimal DOCX files programmatically using python-docx for testing.
"""

import pytest
from docx import Document as DocxDocument

from author_library.errors import ParsingError
from author_library.parsing.docx_parser import DocxParser
from author_library.parsing.models import NodeType


def _create_test_docx(path: object) -> None:
    """Create a minimal DOCX file with headings, paragraphs, and a table."""
    from pathlib import Path

    doc = DocxDocument()
    doc.core_properties.title = "Test DOCX Document"
    doc.core_properties.author = "Test Author"

    doc.add_heading("Chapter 1: Introduction", level=1)
    doc.add_paragraph("This is the first paragraph of the introduction.")
    doc.add_paragraph("This is the second paragraph of the introduction.")

    doc.add_heading("Background", level=2)
    doc.add_paragraph("Background information goes here.")

    doc.add_heading("Chapter 2: Methods", level=1)
    doc.add_paragraph("Description of methods used in this study.")

    # Add a table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header 1"
    table.cell(0, 1).text = "Header 2"
    table.cell(1, 0).text = "Value 1"
    table.cell(1, 1).text = "Value 2"

    doc.save(str(Path(str(path))))


@pytest.fixture
def parser() -> DocxParser:
    return DocxParser()


@pytest.fixture
def test_docx(tmp_path: object) -> object:
    from pathlib import Path

    path = Path(str(tmp_path)) / "test.docx"
    _create_test_docx(path)
    return path


class TestDocxParser:
    async def test_supported_extensions(self, parser: DocxParser) -> None:
        assert parser.supported_extensions() == [".docx"]

    async def test_parse_basic(self, parser: DocxParser, test_docx: object) -> None:
        result = await parser.parse(test_docx)  # type: ignore[arg-type]
        assert result.format == "docx"
        assert result.tree.node_type == NodeType.BOOK

    async def test_metadata_extraction(self, parser: DocxParser, test_docx: object) -> None:
        result = await parser.parse(test_docx)  # type: ignore[arg-type]
        assert result.metadata.title == "Test DOCX Document"
        assert result.metadata.author == "Test Author"

    async def test_chapter_structure(self, parser: DocxParser, test_docx: object) -> None:
        result = await parser.parse(test_docx)  # type: ignore[arg-type]
        chapters = [c for c in result.tree.children if c.node_type == NodeType.CHAPTER]
        assert len(chapters) == 2
        assert chapters[0].metadata.get("title") == "Chapter 1: Introduction"
        assert chapters[1].metadata.get("title") == "Chapter 2: Methods"

    async def test_section_nesting(self, parser: DocxParser, test_docx: object) -> None:
        result = await parser.parse(test_docx)  # type: ignore[arg-type]
        chapters = [c for c in result.tree.children if c.node_type == NodeType.CHAPTER]
        # Chapter 1 should have a "Background" section
        ch1 = chapters[0]
        sections = [c for c in ch1.children if c.node_type == NodeType.SECTION]
        assert len(sections) == 1
        assert sections[0].metadata.get("title") == "Background"

    async def test_paragraph_extraction(self, parser: DocxParser, test_docx: object) -> None:
        result = await parser.parse(test_docx)  # type: ignore[arg-type]

        def find_type(node: object, ntype: NodeType) -> list[object]:
            from author_library.parsing.models import DocumentNode

            assert isinstance(node, DocumentNode)
            found = []
            if node.node_type == ntype:
                found.append(node)
            for child in node.children:
                found.extend(find_type(child, ntype))
            return found

        paragraphs = find_type(result.tree, NodeType.PARAGRAPH)
        assert len(paragraphs) >= 3

    async def test_table_extraction(self, parser: DocxParser, test_docx: object) -> None:
        result = await parser.parse(test_docx)  # type: ignore[arg-type]

        def find_type(node: object, ntype: NodeType) -> list[object]:
            from author_library.parsing.models import DocumentNode

            assert isinstance(node, DocumentNode)
            found = []
            if node.node_type == ntype:
                found.append(node)
            for child in node.children:
                found.extend(find_type(child, ntype))
            return found

        tables = find_type(result.tree, NodeType.TABLE)
        assert len(tables) >= 1
        assert "Header 1" in tables[0].text  # type: ignore[union-attr]

    async def test_raw_text(self, parser: DocxParser, test_docx: object) -> None:
        result = await parser.parse(test_docx)  # type: ignore[arg-type]
        assert "first paragraph" in result.raw_text
        assert result.metadata.word_count > 0

    async def test_file_not_found(self, parser: DocxParser) -> None:
        with pytest.raises(ParsingError, match="not found"):
            await parser.parse("/nonexistent/doc.docx")

    async def test_heading_level(self) -> None:
        assert DocxParser._heading_level("Heading 1") == 1
        assert DocxParser._heading_level("Heading 2") == 2
        assert DocxParser._heading_level("Heading 3") == 3
